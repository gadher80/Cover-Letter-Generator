import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QLineEdit, QPushButton, QWidget, QVBoxLayout, QRadioButton, QTextEdit, QCheckBox
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
import os
import datetime 
from docx import Document
import comtypes.client
#PdfFileMerger
from PyPDF2 import PdfMerger
import os
#import openai
# APIKey = ''
# openai.api_key = APIKey




class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Create labels for the input fields
        self.company_name_label = QLabel("Enter the name of the company:", self)
        self.company_address_label = QLabel("Enter the address of the company with this format: (Berliner Straße 70, 70777, Berlin)", self)
        self.position_name_label = QLabel("Enter the name of the position you're applying to:", self)
        self.domain_label = QLabel("Enter the domain what you want to add in introduction paragraph:", self)
        self.position_code_label = QLabel("Enter the position code (if needed):", self)
        self.particular_person_label = QLabel("Enter the name of the particular person (if needed):", self)
        self.cover_letter_label = QLabel("Enter or edit the cover letter:", self)
        # Create line edits for the input fields
        self.company_name_edit = QLineEdit(self)
        self.company_address_edit = QLineEdit(self)
        self.position_name_edit = QLineEdit(self)
        #add radio button for domain
        self.data_science_radio = QRadioButton("Data Science", self)
        self.machine_learning_radio = QRadioButton("Machine Learning", self)
        self.data_analytics_radio = QRadioButton("Data Analytics", self)
        self.software_development_radio = QRadioButton("Software Development", self)
        self.position_code_edit = QLineEdit(self)
        self.particular_person_edit = QLineEdit(self)
        #add check box if you want to attach resume 
        self.attach_resume = QCheckBox("Attach Resume", self)
        #add check box if you want to attach cover letter   
        #craete a text edit wheere I can edit the text
        self.cover_letter = QTextEdit(self)
        #also add defualt text to the text edit

        # Create a button to run the Python script
        self.run_button = QPushButton("Create", self)
        self.run_button.clicked.connect(self.run_script)

        # Set the layout for the UI
        central_widget = QWidget()
        layout = QVBoxLayout(central_widget)
        layout.addWidget(self.company_name_label)
        layout.addWidget(self.company_name_edit)
        layout.addWidget(self.company_address_label)
        layout.addWidget(self.company_address_edit)
        layout.addWidget(self.position_name_label)
        layout.addWidget(self.position_name_edit)
        layout.addWidget(self.domain_label)
        layout.addWidget(self.data_science_radio)
        layout.addWidget(self.machine_learning_radio)
        layout.addWidget(self.data_analytics_radio)
        layout.addWidget(self.software_development_radio)

        layout.addWidget(self.position_code_label)
        layout.addWidget(self.position_code_edit)
        layout.addWidget(self.particular_person_label)
        layout.addWidget(self.particular_person_edit)
        
        layout.addWidget(self.cover_letter_label)
        layout.addWidget(self.cover_letter)

        #add this check box left side of the run button
        layout.addWidget(self.attach_resume)
        layout.addWidget(self.run_button)
        self.setCentralWidget(central_widget)
        #connect the radio button to the text edit
        self.data_science_radio.clicked.connect(self.update_texBox)
        self.machine_learning_radio.clicked.connect(self.update_texBox)
        self.data_analytics_radio.clicked.connect(self.update_texBox)
        self.software_development_radio.clicked.connect(self.update_texBox)

        #update the textBox if particular_person_edit is changed
        self.particular_person_edit.textChanged.connect(self.update_texBox)

    #write function resume check box is checked
                    

        

    #change the cover_letter if radio button is clicked and change the text
    def update_texBox(self):
        global domain
        domain = ''
        if particular_person_edit := self.particular_person_edit.text():
            if particular_person_edit.startswith('Herr'):
                particular_person_edit = (
                    f'Sehr geehrter {self.particular_person_edit.text()}'
                )
            else: 
                particular_person_edit = (
                    f'Sehr geehrte {self.particular_person_edit.text()}'
                )
        else:
            particular_person_edit = 'Sehr geehrte Damen und Herren'

        # import openai
        # openai.api_key ="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"
        # prompt = f' {self.position_name_edit.text()} cover letter 200 words at {self.company_name_edit.text()}  to {particular_person_edit}'
        # model = "text-curie-001"

        # response = openai.Completion.create(
        #     engine=model,
        #     prompt=prompt,
        #     max_tokens=2000,
        #     n=1,
        #     stop=None,
        #     temperature=0,
        # )

        # #print(response.choices[0].text)
        # global displaytext
        # displaytext2 = response.choices[0].text
        # print(displaytext2)
        displaytext = f"{particular_person_edit},\n\nI am writing to apply for the position of {domain} at your company. I am a recent graduate of the University of Berlin with a degree in Computer Science and a minor in Mathematics. I am interested in the position because I am passionate about data science and I believe that I have the skills and experience to be a valuable asset to your company.\n\nI have a strong background in data science, including machine learning, data analytics, and software development. I have experience with Python, R, and SQL. I have also worked on several projects that involve data science, including a project that involved using machine learning to predict the price of a house based on its features. I have also worked on a project that involved using data analytics to analyze the data of a company and make recommendations for the company.\n\nI am excited about the opportunity to work at your company and I look forward to hearing from you.\n\nSincerely,\n\nHardik Gadher"

        if self.data_science_radio.isChecked():
            domain = 'Data Scientist'
            
        elif self.machine_learning_radio.isChecked():
            domain = 'Machine Learning Engineer'
            
        elif self.data_analytics_radio.isChecked():
            domain = 'Data Analyst'
            
        elif self.software_development_radio.isChecked():
            domain = 'Software Developer'
        self.cover_letter.setText(displaytext)
   

    def run_script(self):  # sourcery skip: low-code-quality

        company_name = self.company_name_edit.text()
        company_address = self.company_address_edit.text()
        position_name = self.position_name_edit.text()
        position_code = self.position_code_edit.text()
        address_list = company_address.split(", ")
        document = Document('coverletter.docx')
        tables = document.tables
        tables[1].cell(0,0).text = company_name
        tables[1].cell(0,1).text = address_list[0]
        tables[1].cell(0,2).text = f'{address_list[1]} {address_list[2]}'
        #change font size of the company name
        tables[1].cell(0,0).paragraphs[0].runs[0].font.size = Pt(14)
        tables[1].cell(0,0).paragraphs[0].runs[0].font.bold = True
        tables[1].cell(0,0).paragraphs[0].runs[0].font.name = 'Arial'
        #change font size for the company address and color to #737373
        tables[1].cell(0,1).paragraphs[0].runs[0].font.size = Pt(12)
        tables[1].cell(0,1).paragraphs[0].runs[0].font.name = 'Arial'
        tables[1].cell(0,2).paragraphs[0].runs[0].font.size = Pt(12)
        tables[1].cell(0,2).paragraphs[0].runs[0].font.name = 'Arial'
        tables[1].cell(0,1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0x73, 0x73, 0x73)
        tables[1].cell(0,2).paragraphs[0].runs[0].font.color.rgb = RGBColor(0x73, 0x73, 0x73)
        tables[3].cell(0,0).text = f"Bewerbung als {position_name}"
        #change font size of the position name
        tables[3].cell(0,0).paragraphs[0].runs[0].font.size = Pt(14)
        tables[3].cell(0,0).paragraphs[0].runs[0].font.bold = True
        tables[3].cell(0,0).paragraphs[0].runs[0].font.name = 'Arial'
        if position_code:
            tables[3].add_row()
            tables[3].cell(0,1).text = f"Kennziffer: {position_code}"
            tables[3].cell(0,1).paragraphs[0].runs[0].font.size = Pt(14)
            tables[3].cell(0,1).paragraphs[0].runs[0].font.name = 'Arial'
            tables[3].cell(0,1).paragraphs[0].runs[0].font.bold = True
        
        now = datetime.datetime.now()
        months = ['Januar', 'Februar', 'März', 'April', 'Mai', 'Juni', 'Juli', 'August', 'September', 'Oktober', 'November', 'Dezember']
        date_string = "{0:%d}. {1} {0:%Y}".format(now, months[now.month - 1]) 
        tables[2].cell(0,0).text = f"Erfurt {date_string}"
        #justify right
        tables[2].cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #make bold
        tables[2].cell(0,0).paragraphs[0].runs[0].font.bold = True
        tables[2].cell(0,0).paragraphs[0].runs[0].font.size = Pt(12)
        tables[2].cell(0,0).paragraphs[0].runs[0].font.name = 'Arial'
        tables[4].cell(0,0).text = self.cover_letter.toPlainText()
        tables[4].cell(0,0).paragraphs[0].runs[0].font.size = Pt(12)
        tables[4].cell(0,0).paragraphs[0].runs[0].font.name = 'Arial'
        tables[4].cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        tables[4].cell(0,0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0x73, 0x73, 0x73)
        document.save(f"{company_name} {position_name} {now.strftime('%d-%m-%Y')}.docx")
        #save as pdf
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(f"{company_name} {position_name} {now.strftime('%d-%m-%Y')}.docx"))
        doc.SaveAs(os.path.abspath(f"{company_name} {position_name} {now.strftime('%d-%m-%Y')}.pdf"), FileFormat=17)
        doc.Close()
        word.Quit()

        #if resume check box is checked attached default resume
        if self.attach_resume.isChecked():
            merger = PdfMerger()
            pdf_files = [f"{company_name} {position_name} {now.strftime('%d-%m-%Y')}.pdf", "CV.pdf"]
            for pdf_file in pdf_files:
                #Append PDF files
                merger.append(pdf_file)
            merger.write(f"{company_name} {position_name} {now.strftime('%d-%m-%Y')}.pdf")
            merger.close()

        os.startfile(f"{company_name} {position_name} {datetime.datetime.now().strftime('%d-%m-%Y')}.pdf")
       



if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
